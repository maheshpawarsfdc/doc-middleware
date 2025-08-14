from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
import base64
import pdfplumber
from docx import Document
from io import BytesIO
import requests
import os
import logging
from typing import List, Optional
import time
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
import asyncio
from contextlib import asynccontextmanager
import json
from typing import Dict, Any

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Application lifespan manager
@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup
    logger.info("üöÄ FastAPI server starting up...")
    
    # Validate configuration
    if not os.environ.get('GEMINI_API_KEY'):
        logger.error("‚ùå GEMINI_API_KEY environment variable not set!")
        raise RuntimeError("GEMINI_API_KEY is required")
    
    logger.info("‚úÖ Configuration validated")
    yield
    
    # Shutdown
    logger.info("üõë FastAPI server shutting down...")

app = FastAPI(
    title="Document Analysis Middleware", 
    version="2.0.0",
    lifespan=lifespan
)

# Configure requests session with retry strategy
def create_requests_session():
    session = requests.Session()
    retry_strategy = Retry(
        total=3,
        backoff_factor=2,
        status_forcelist=[429, 500, 502, 503, 504],
        respect_retry_after_header=True
    )
    adapter = HTTPAdapter(max_retries=retry_strategy)
    session.mount("http://", adapter)
    session.mount("https://", adapter)
    return session

requests_session = create_requests_session()

# ----------------------------
# Models
# ----------------------------
class FileInput(BaseModel):
    filename: str
    filedata: str
    customPrompt: Optional[str] = None

class Message(BaseModel):
    role: str
    content: str

class FollowupPayload(BaseModel):
    messages: List[Message]

# ----------------------------
# Configuration
# ----------------------------
MAX_TEXT_LENGTH = 80000
GEMINI_TIMEOUT = 90
MAX_FILE_SIZE = 15 * 1024 * 1024
MAX_RETRIES = 3
INITIAL_RETRY_DELAY = 2

# ----------------------------
# Error Handlers
# ----------------------------
@app.exception_handler(HTTPException)
async def http_exception_handler(request, exc):
    logger.error(f"HTTP Exception: {exc.status_code} - {exc.detail}")
    return {
        "error": exc.detail,
        "status_code": exc.status_code,
        "timestamp": time.time()
    }

@app.exception_handler(Exception)
async def general_exception_handler(request, exc):
    logger.error(f"Unexpected error: {str(exc)}", exc_info=True)
    return {
        "error": "Internal server error occurred",
        "status_code": 500,
        "timestamp": time.time()
    }

# ----------------------------
# Document Type Detection
# ----------------------------
def detect_document_type(extracted_text: str, filename: str) -> str:
    """Detect document type based on content and filename"""
    text_lower = extracted_text.lower()
    filename_lower = filename.lower()
    
    # Resume detection
    resume_keywords = ['resume', 'cv', 'curriculum vitae', 'experience', 'education', 'skills', 'objective', 'summary']
    if any(keyword in filename_lower for keyword in ['resume', 'cv']) or \
       sum(1 for keyword in resume_keywords if keyword in text_lower) >= 3:
        return 'RESUME'
    
    # Contract detection
    contract_keywords = ['agreement', 'contract', 'terms', 'conditions', 'party', 'obligations', 'consideration']
    if any(keyword in filename_lower for keyword in ['contract', 'agreement']) or \
       sum(1 for keyword in contract_keywords if keyword in text_lower) >= 4:
        return 'CONTRACT'
    
    # NDA detection
    nda_keywords = ['confidential', 'non-disclosure', 'proprietary', 'trade secret', 'confidentiality']
    if any(keyword in filename_lower for keyword in ['nda', 'confidential']) or \
       sum(1 for keyword in nda_keywords if keyword in text_lower) >= 2:
        return 'NDA'
    
    # Policy detection
    policy_keywords = ['policy', 'procedure', 'guidelines', 'rules', 'regulations', 'compliance']
    if any(keyword in filename_lower for keyword in ['policy', 'procedure']) or \
       sum(1 for keyword in policy_keywords if keyword in text_lower) >= 3:
        return 'POLICY'
    
    # Offer Letter detection
    offer_keywords = ['offer', 'position', 'salary', 'compensation', 'start date', 'employment']
    if any(keyword in filename_lower for keyword in ['offer']) or \
       sum(1 for keyword in offer_keywords if keyword in text_lower) >= 4:
        return 'OFFER_LETTER'
    
    return 'GENERAL_DOCUMENT'

# ----------------------------
# /process Endpoint
# ----------------------------
@app.post("/process")
async def process_file(input: FileInput):
    request_id = str(int(time.time() * 1000))
    logger.info(f"[{request_id}] üì© Processing file: {input.filename}")
    
    try:
        # Input validation
        if not input.filename:
            raise HTTPException(status_code=400, detail="Filename is required")
        
        if not input.filedata:
            raise HTTPException(status_code=400, detail="File data is required")
        
        # Decode and validate file data
        try:
            binary = base64.b64decode(input.filedata)
        except Exception as e:
            logger.error(f"[{request_id}] ‚ùå Base64 decode error: {str(e)}")
            raise HTTPException(status_code=400, detail="Invalid base64 file data")
        
        # Check file size
        if len(binary) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=413, 
                detail=f"File too large. Maximum size: {MAX_FILE_SIZE/1024/1024:.1f}MB"
            )
        
        logger.info(f"[{request_id}] üìä File size: {len(binary)} bytes")
        
        # Extract text based on file type
        extracted_text = await extract_text_from_file(input.filename, binary, request_id)
        
        if not extracted_text.strip():
            raise HTTPException(status_code=400, detail="No text could be extracted from the document")

        # Truncate if too long
        if len(extracted_text) > MAX_TEXT_LENGTH:
            logger.warning(f"[{request_id}] ‚ö†Ô∏è Text truncated from {len(extracted_text)} to {MAX_TEXT_LENGTH} characters")
            extracted_text = extracted_text[:MAX_TEXT_LENGTH] + "\n\n[Document truncated due to length...]"
        
        logger.info(f"[{request_id}] üìÑ Extracted {len(extracted_text)} characters of text")

        # Detect document type
        doc_type = detect_document_type(extracted_text, input.filename)
        logger.info(f"[{request_id}] üîç Detected document type: {doc_type}")

        # Build the ENHANCED prompt based on document type
        base_prompt = build_enhanced_analysis_prompt(extracted_text, doc_type)
        final_prompt = (
            input.customPrompt.strip() + "\n\n" + base_prompt
            if input.customPrompt else base_prompt
        )
        
        # Call Gemini API with proper error handling and retry
        response = await call_gemini_with_retry(final_prompt, request_id)
        
        logger.info(f"[{request_id}] ‚úÖ Analysis completed successfully")
        return {
            "insights": response, 
            "status": "success",
            "request_id": request_id,
            "timestamp": time.time(),
            "document_type": doc_type
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[{request_id}] ‚ùå Unexpected error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Processing failed: {str(e)}")

# ----------------------------
# Text Extraction Function
# ----------------------------
async def extract_text_from_file(filename: str, binary: bytes, request_id: str) -> str:
    """Extract text from PDF or DOCX files with improved error handling"""
    extracted_text = ""
    
    try:
        if filename.lower().endswith(".pdf"):
            logger.info(f"[{request_id}] üìÑ Processing PDF file")
            
            try:
                with pdfplumber.open(BytesIO(binary)) as pdf:
                    total_pages = len(pdf.pages)
                    logger.info(f"[{request_id}] üìä PDF has {total_pages} pages")
                    
                    if total_pages == 0:
                        raise ValueError("PDF file appears to be empty or corrupted")
                    
                    for i, page in enumerate(pdf.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                extracted_text += page_text + "\n"
                            logger.debug(f"[{request_id}] ‚úÖ Processed page {i+1}/{total_pages}")
                        except Exception as e:
                            logger.warning(f"[{request_id}] ‚ö†Ô∏è Error extracting page {i+1}: {str(e)}")
                            continue
            except Exception as e:
                logger.error(f"[{request_id}] ‚ùå PDF processing failed: {str(e)}")
                raise HTTPException(status_code=422, detail=f"PDF processing failed: {str(e)}")
                        
        elif filename.lower().endswith(".docx"):
            logger.info(f"[{request_id}] üìÑ Processing DOCX file")
            
            try:
                doc = Document(BytesIO(binary))
                total_paragraphs = len(doc.paragraphs)
                logger.info(f"[{request_id}] üìä DOCX has {total_paragraphs} paragraphs")
                
                if total_paragraphs == 0:
                    raise ValueError("DOCX file appears to be empty or corrupted")
                
                for i, para in enumerate(doc.paragraphs):
                    try:
                        if para.text.strip():
                            extracted_text += para.text + "\n"
                    except Exception as e:
                        logger.warning(f"[{request_id}] ‚ö†Ô∏è Error extracting paragraph {i+1}: {str(e)}")
                        continue
            except Exception as e:
                logger.error(f"[{request_id}] ‚ùå DOCX processing failed: {str(e)}")
                raise HTTPException(status_code=422, detail=f"DOCX processing failed: {str(e)}")
        else:
            raise HTTPException(
                status_code=400, 
                detail="Unsupported file type. Only PDF and DOCX are supported."
            )
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[{request_id}] ‚ùå Text extraction failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Text extraction failed: {str(e)}")
    
    return extracted_text.strip()

# ----------------------------
# ENHANCED Document-Type-Specific Prompt Building
# ----------------------------
def build_enhanced_analysis_prompt(extracted_text: str, doc_type: str) -> str:
    """Build document-type specific analysis prompts"""
    
    # Common formatting rules
    formatting_rules = """
**CRITICAL FORMATTING RULES - MUST BE FOLLOWED EXACTLY:**

1. **FOR DATES:** You MUST provide context for EVERY date. NEVER list raw dates.
   ‚ùå WRONG: "- January 15, 2024"
   ‚úÖ CORRECT: "- Employment Start Date: January 15, 2024"

2. **FOR MONETARY VALUES:** You MUST provide context for EVERY amount. NEVER list raw amounts.
   ‚ùå WRONG: "- $185,000"
   ‚úÖ CORRECT: "- Annual Salary: $185,000"

3. **IF CONTEXT IS UNCLEAR:** Use descriptive labels based on document context.
"""

    if doc_type == 'RESUME':
        return f"""
You are an expert HR recruitment specialist analyzing resumes for talent acquisition.

{formatting_rules}

**RESUME ANALYSIS INSTRUCTIONS:**
Analyze this resume for an HR team evaluating candidate fit, experience relevance, and potential red flags.

**OUTPUT FORMAT - FOLLOW EXACTLY:**

**Document Type & Classification**
Resume/CV Analysis

**Document Summary (MANDATORY ‚Äî 10‚Äì12 sentences)**
Provide a comprehensive candidate overview that eliminates the need to read the full resume.

Your summary MUST clearly state:
1. Candidate's name, current/most recent role, and years of experience
2. Educational background and key qualifications
3. Core technical skills and competencies
4. Industry experience and domain expertise
5. Career progression and growth trajectory
6. Notable achievements, certifications, or awards
7. Employment gaps or career transitions (if any)
8. Salary expectations or current compensation (if mentioned)
9. Location, visa status, or relocation requirements
10. Overall candidate strength assessment and recommended next steps

**Key Information Extracted**

**Candidate Profile:**
- Full Name: [Name]
- Current Role: [Title] at [Company]
- Total Experience: [X years in Y domain]
- Location: [City, State/Country]
- Contact: [Email, Phone if provided]

**Education & Certifications:**
- Highest Degree: [Degree] from [Institution] - [Graduation Year]
- Professional Certifications: [List with dates]
- Relevant Training: [Technical or professional development]

**Technical Skills:**
- Programming Languages: [If applicable]
- Software/Tools: [Professional tools, platforms]
- Industry Knowledge: [Domain-specific expertise]

**Experience Timeline:**
[For each role, format as:]
- Role Duration: [Start Date] to [End Date] ([X months/years])
- Position: [Job Title] at [Company Name]
- Key Responsibilities: [Brief summary]

**Employment Analysis:**
- Career Progression: [Upward/lateral/mixed]
- Employment Gaps: [Any gaps > 6 months with explanation if provided]
- Job Stability: [Average tenure per role]
- Industry Focus: [Consistent/diverse sectors]

**Compensation Information:**
- Current/Expected Salary: [If mentioned]
- Previous Compensation: [If historical data available]
- Benefits Expectations: [If specified]

**Compliance & Risk Assessment**

**Potential Red Flags:**
- [HIGH/MEDIUM/LOW] Employment gaps without explanation
- [HIGH/MEDIUM/LOW] Frequent job changes (< 1 year tenure)
- [HIGH/MEDIUM/LOW] Inconsistent career progression
- [HIGH/MEDIUM/LOW] Missing contact information
- [HIGH/MEDIUM/LOW] Overqualification for target role
- [HIGH/MEDIUM/LOW] Skills mismatch with requirements

**Missing or Unclear Elements:**
- [Items typically expected on resumes but absent]
- [Incomplete employment dates]
- [Vague job descriptions or achievements]

**Verification Requirements:**
- [Education verification needed]
- [Employment verification priorities]
- [Professional references status]

**Actionable Recommendations**

**Interview Focus Areas:**
- [Key technical areas to assess]
- [Behavioral questions to explore]
- [Career transition explanations needed]

**Next Steps:**
- [Recommend phone screen/technical interview/skip]
- [Salary negotiation considerations]
- [Reference check priorities]
- [Additional documentation needed]

**Stakeholder Actions:**
- [Hiring manager review requirements]
- [Technical team assessment needs]
- [HR follow-up items]

**DOCUMENT CONTENT TO ANALYZE:**
{extracted_text}
"""

    elif doc_type == 'CONTRACT':
        return f"""
You are an expert legal and business contracts analyst specializing in commercial agreements.

{formatting_rules}

**CONTRACT ANALYSIS INSTRUCTIONS:**
Analyze this contract for a legal/business team reviewing terms, obligations, risks, and compliance requirements.

**OUTPUT FORMAT - FOLLOW EXACTLY:**

**Document Type & Classification**
[Contract Type: Service Agreement/Purchase Agreement/Employment Contract/etc.]

**Document Summary (MANDATORY ‚Äî 10‚Äì12 sentences)**
Provide a comprehensive contract overview for executive review.

Your summary MUST clearly state:
1. Contract type, purpose, and primary business objective
2. All contracting parties with their roles and legal entities
3. Contract value, payment terms, and financial obligations
4. Key deliverables, services, or products covered
5. Contract duration, start/end dates, and renewal terms
6. Critical deadlines, milestones, or performance requirements
7. Liability limitations, indemnification, and risk allocation
8. Termination conditions and notice requirements
9. Governing law, dispute resolution mechanisms
10. Overall risk assessment and recommended actions

**Key Information Extracted**

**Contracting Parties:**
- Primary Party: [Company Name] - [Role: Buyer/Seller/Service Provider]
- Secondary Party: [Company Name] - [Role and legal status]
- Authorized Signatories: [Names and titles]

**Financial Terms:**
- Contract Value: [Total amount with currency]
- Payment Schedule: [Due dates and amounts]
- Late Payment Penalties: [Fees and interest rates]
- Expense Responsibilities: [Who pays what]

**Performance Obligations:**
- Primary Deliverables: [Services, products, or outcomes]
- Quality Standards: [Acceptance criteria and SLAs]
- Timeline Requirements: [Key milestones and deadlines]
- Performance Metrics: [How success is measured]

**Important Dates:**
- Contract Effective Date: [Start date]
- Performance Period: [Duration of obligations]
- Key Milestone Dates: [Critical deadlines]
- Contract Expiration: [End date]
- Renewal Deadline: [Notice required by date]

**Risk & Liability Terms:**
- Liability Caps: [Maximum exposure amounts]
- Indemnification Scope: [What's covered]
- Insurance Requirements: [Types and amounts]
- Force Majeure Provisions: [Covered events]

**Compliance & Risk Assessment**

**Contract Risks:**
- [HIGH/MEDIUM/LOW] Unlimited liability exposure
- [HIGH/MEDIUM/LOW] Aggressive penalty terms
- [HIGH/MEDIUM/LOW] Unclear deliverable definitions
- [HIGH/MEDIUM/LOW] Insufficient termination protection
- [HIGH/MEDIUM/LOW] Missing governing law clauses
- [HIGH/MEDIUM/LOW] Inadequate intellectual property terms

**Missing Critical Clauses:**
- [Standard clauses that should be present but aren't]
- [Industry-specific requirements not addressed]
- [Regulatory compliance terms missing]

**Regulatory Considerations:**
- [Industry regulations that apply]
- [Compliance requirements identified]
- [Data protection and privacy obligations]

**Actionable Recommendations**

**Immediate Actions Required:**
- [Contract terms requiring negotiation]
- [Legal review priorities]
- [Approval workflow requirements]

**Risk Mitigation Steps:**
- [Recommended contract amendments]
- [Additional protections needed]
- [Insurance verification requirements]

**Ongoing Management:**
- [Performance monitoring requirements]
- [Renewal timeline and considerations]
- [Compliance tracking needs]

**DOCUMENT CONTENT TO ANALYZE:**
{extracted_text}
"""

    elif doc_type == 'NDA':
        return f"""
You are an expert legal counsel specializing in confidentiality and non-disclosure agreements.

{formatting_rules}

**NDA ANALYSIS INSTRUCTIONS:**
Analyze this NDA for a legal team reviewing confidentiality terms, scope, and enforceability.

**OUTPUT FORMAT - FOLLOW EXACTLY:**

**Document Type & Classification**
Non-Disclosure Agreement (NDA) / Confidentiality Agreement

**Document Summary (MANDATORY ‚Äî 10‚Äì12 sentences)**
Provide a comprehensive NDA overview for legal and business teams.

Your summary MUST clearly state:
1. NDA type (mutual/unilateral) and primary business context
2. All parties bound by confidentiality obligations
3. Definition and scope of confidential information covered
4. Permitted uses and restrictions on confidential information
5. Duration of confidentiality obligations and survival terms
6. Geographic scope and applicable jurisdictions
7. Exceptions to confidentiality (public domain, prior knowledge, etc.)
8. Return or destruction requirements for confidential materials
9. Remedies for breach including injunctive relief provisions
10. Overall enforceability assessment and risk evaluation

**Key Information Extracted**

**NDA Parties:**
- Disclosing Party: [Company/Individual Name and Role]
- Receiving Party: [Company/Individual Name and Role]
- NDA Type: [Mutual/Unilateral/Multilateral]

**Confidentiality Scope:**
- Information Definition: [What constitutes confidential information]
- Marking Requirements: [How confidential info must be identified]
- Oral Information: [How verbal disclosures are handled]
- Third Party Information: [Treatment of others' confidential data]

**Permitted Uses:**
- Authorized Purposes: [What receiving party can do with info]
- Internal Distribution: [Who within organization can access]
- Employee Obligations: [How employees are bound]

**Important Dates:**
- NDA Effective Date: [When obligations begin]
- Confidentiality Period: [How long obligations last]
- Information Return Deadline: [When materials must be returned]
- NDA Expiration: [When agreement terminates]

**Restrictions & Obligations:**
- Use Limitations: [What receiving party cannot do]
- Security Requirements: [How information must be protected]
- Disclosure Restrictions: [Who information cannot be shared with]
- Notification Requirements: [Breach reporting obligations]

**Compliance & Risk Assessment**

**NDA Risks:**
- [HIGH/MEDIUM/LOW] Overly broad confidentiality definition
- [HIGH/MEDIUM/LOW] Indefinite confidentiality period
- [HIGH/MEDIUM/LOW] Unclear permitted use scope
- [HIGH/MEDIUM/LOW] Weak enforcement mechanisms
- [HIGH/MEDIUM/LOW] Missing return/destruction requirements
- [HIGH/MEDIUM/LOW] Inadequate exception clauses

**Enforceability Issues:**
- [Geographic scope enforceability concerns]
- [Duration reasonableness assessment]
- [Definition clarity and specificity]
- [Consideration adequacy]

**Legal Compliance:**
- [Jurisdiction-specific requirements]
- [Industry regulation considerations]
- [International data transfer implications]

**Actionable Recommendations**

**Immediate Legal Actions:**
- [Terms requiring clarification or amendment]
- [Missing standard NDA provisions to add]
- [Negotiation priorities for business protection]

**Risk Management:**
- [Information handling process improvements]
- [Employee training requirements]
- [Monitoring and compliance procedures]

**Business Operations:**
- [Information sharing guidelines needed]
- [Documentation and tracking systems]
- [Breach response procedures]

**DOCUMENT CONTENT TO ANALYZE:**
{extracted_text}
"""

    elif doc_type == 'OFFER_LETTER':
        return f"""
You are an expert HR specialist analyzing employment offer letters for completeness and compliance.

{formatting_rules}

**OFFER LETTER ANALYSIS INSTRUCTIONS:**
Analyze this offer letter for HR teams reviewing employment terms, compliance, and completeness.

**OUTPUT FORMAT - FOLLOW EXACTLY:**

**Document Type & Classification**
Employment Offer Letter

**Document Summary (MANDATORY ‚Äî 10‚Äì12 sentences)**
Provide a comprehensive offer analysis for HR and hiring managers.

Your summary MUST clearly state:
1. Candidate name, offered position, and reporting structure
2. Compensation package details (salary, bonus, equity, benefits)
3. Employment start date and location requirements
4. Employment type (full-time/part-time, exempt/non-exempt status)
5. Key terms and conditions of employment
6. Probationary period and performance evaluation timeline
7. Confidentiality, non-compete, or other restrictive covenants
8. Offer acceptance deadline and next steps required
9. Benefits eligibility and enrollment information
10. Compliance with employment laws and company policies

**Key Information Extracted**

**Candidate & Position:**
- Candidate Name: [Full name]
- Position Title: [Job title and level]
- Department: [Organizational unit]
- Reporting Manager: [Direct supervisor name and title]
- Employment Type: [Full-time/Part-time, Exempt/Non-exempt]

**Compensation Package:**
- Base Salary: [Annual/hourly rate]
- Bonus Eligibility: [Performance/signing bonus details]
- Equity Compensation: [Stock options, RSUs, vesting schedule]
- Commission Structure: [If applicable]

**Important Dates:**
- Offer Issue Date: [When offer was extended]
- Response Deadline: [Acceptance deadline]
- Employment Start Date: [First day of work]
- Probation Period End: [If applicable]
- First Performance Review: [Scheduled evaluation date]

**Benefits & Perquisites:**
- Health Insurance: [Coverage effective date and options]
- Retirement Benefits: [401k, pension, vesting details]
- Paid Time Off: [Vacation, sick leave, personal days]
- Professional Development: [Training, education allowances]
- Other Benefits: [Specific perks or allowances]

**Employment Terms:**
- Work Location: [Office, remote, hybrid requirements]
- Work Schedule: [Hours, flexibility arrangements]
- Travel Requirements: [Expected travel percentage]
- At-Will Employment: [Employment relationship terms]

**Compliance & Risk Assessment**

**Employment Law Compliance:**
- [HIGH/MEDIUM/LOW] Fair Labor Standards Act compliance
- [HIGH/MEDIUM/LOW] Equal employment opportunity compliance
- [HIGH/MEDIUM/LOW] State wage and hour law adherence
- [HIGH/MEDIUM/LOW] Immigration compliance (I-9 requirements)
- [HIGH/MEDIUM/LOW] Background check and drug testing clarity

**Missing Standard Elements:**
- [Required disclosures not present]
- [Standard benefits not mentioned]
- [Typical employment terms absent]
- [Legal compliance statements missing]

**Risk Areas:**
- [Unclear compensation terms]
- [Missing confidentiality agreements]
- [Incomplete benefits information]
- [Vague performance expectations]

**Actionable Recommendations**

**Before Candidate Acceptance:**
- [Clarifications needed in offer terms]
- [Additional documentation to provide]
- [Legal review requirements]
- [Approval workflow completion]

**Post-Acceptance Actions:**
- [Onboarding documentation needed]
- [Background check and verification steps]
- [Benefits enrollment coordination]
- [Equipment and access provisioning]

**HR Process Improvements:**
- [Template updates recommended]
- [Compliance verification steps]
- [Manager communication requirements]
- [Documentation and filing needs]

**DOCUMENT CONTENT TO ANALYZE:**
{extracted_text}
"""

    elif doc_type == 'POLICY':
        return f"""
You are an expert compliance and policy analyst reviewing organizational policies and procedures.

{formatting_rules}

**POLICY ANALYSIS INSTRUCTIONS:**
Analyze this policy document for a compliance team reviewing policy effectiveness, gaps, and regulatory alignment.

**OUTPUT FORMAT - FOLLOW EXACTLY:**

**Document Type & Classification**
[Policy Type: HR Policy/IT Policy/Safety Policy/Compliance Policy/etc.]

**Document Summary (MANDATORY ‚Äî 10‚Äì12 sentences)**
Provide a comprehensive policy overview for compliance and management teams.

Your summary MUST clearly state:
1. Policy name, scope, and primary regulatory/business purpose
2. Affected employee populations and organizational units
3. Key policy requirements and mandatory procedures
4. Prohibited behaviors and activities clearly defined
5. Compliance monitoring and enforcement mechanisms
6. Training requirements and awareness programs
7. Reporting procedures and escalation paths
8. Policy review schedule and update responsibilities
9. Regulatory alignment and legal compliance status
10. Implementation gaps and recommended improvements

**Key Information Extracted**

**Policy Details:**
- Policy Title: [Official policy name]
- Policy Number: [Document identifier if present]
- Effective Date: [When policy becomes active]
- Review Schedule: [How often policy is reviewed]
- Policy Owner: [Department or individual responsible]

**Scope & Coverage:**
- Applicable Employees: [Who must follow this policy]
- Geographic Scope: [Locations where policy applies]
- Exceptions: [Any groups or situations excluded]
- Related Policies: [Connected or dependent policies]

**Key Requirements:**
- Mandatory Actions: [What employees must do]
- Prohibited Activities: [What is not allowed]
- Approval Requirements: [Who must authorize what]
- Documentation Standards: [Required record keeping]

**Important Dates:**
- Policy Effective Date: [Implementation date]
- Next Review Date: [Scheduled policy review]
- Training Completion Deadline: [Employee training requirements]
- Compliance Reporting Due: [Regular reporting schedule]

**Compliance Framework:**
- Regulatory Basis: [Laws, regulations, or standards addressed]
- Monitoring Procedures: [How compliance is measured]
- Audit Requirements: [Internal/external audit procedures]
- Violation Consequences: [Disciplinary actions for non-compliance]

**Compliance & Risk Assessment**

**Policy Gaps & Risks:**
- [HIGH/MEDIUM/LOW] Unclear procedure definitions
- [HIGH/MEDIUM/LOW] Missing enforcement mechanisms
- [HIGH/MEDIUM/LOW] Inadequate training requirements
- [HIGH/MEDIUM/LOW] Weak monitoring procedures
- [HIGH/MEDIUM/LOW] Outdated regulatory references
- [HIGH/MEDIUM/LOW] Insufficient escalation procedures

**Regulatory Alignment:**
- [Current regulatory requirements covered]
- [Compliance gaps identified]
- [Industry best practices comparison]
- [Legal risk assessment]

**Implementation Challenges:**
- [Resource requirements for compliance]
- [Training and communication needs]
- [Technology or system dependencies]
- [Cultural or behavioral change requirements]

**Actionable Recommendations**

**Immediate Policy Updates:**
- [Critical gaps requiring immediate attention]
- [Regulatory compliance updates needed]
- [Clarifications required for better understanding]

**Implementation Improvements:**
- [Enhanced training program development]
- [Better monitoring and reporting systems]
- [Improved communication strategies]
- [Technology solutions for compliance tracking]

**Ongoing Management:**
- [Regular review and update procedures]
- [Performance metrics and KPIs to track]
- [Stakeholder engagement requirements]
- [Continuous improvement processes]

**DOCUMENT CONTENT TO ANALYZE:**
{extracted_text}
"""

    else:  # GENERAL_DOCUMENT
        return f"""
You are an expert document analysis assistant specialized in business document review.

{formatting_rules}

**GENERAL DOCUMENT ANALYSIS INSTRUCTIONS:**
Analyze this document for a business team reviewing content for key information, risks, and actionable insights.

**OUTPUT FORMAT - FOLLOW EXACTLY:**

**Document Type & Classification**
[Document Type: Report/Memo/Proposal/Agreement/Letter/etc.]

**Document Summary (MANDATORY ‚Äî 10‚Äì12 sentences)**
Provide a comprehensive document overview that eliminates the need to read the full document.

Your summary MUST clearly state:
1. Document type, purpose, and primary business objective
2. Key parties, stakeholders, or entities involved
3. Main topics, issues, or subjects addressed
4. Important decisions, recommendations, or proposals
5. Financial information, costs, or budget implications
6. Timeline, deadlines, or critical dates mentioned
7. Risks, concerns, or potential issues identified
8. Required actions, approvals, or follow-up steps
9. Compliance, regulatory, or legal considerations
10. Overall significance and recommended next steps

**Key Information Extracted**

**Document Overview:**
- Document Title: [Title or subject]
- Document Date: [Creation or effective date]
- Author/Sender: [Who created or sent the document]
- Recipients: [Intended audience or recipients]
- Document Purpose: [Primary objective or reason]

**Key Stakeholders:**
- Primary Parties: [Main individuals or organizations involved]
- Decision Makers: [Who has authority or approval rights]
- Affected Parties: [Who is impacted by the document content]

**Important Dates:**
- Document Date: [When document was created]
- Effective Date: [When content becomes active]
- Deadline Dates: [Critical timelines or due dates]
- Review Dates: [Scheduled follow-up or assessment dates]

**Financial Information:**
- Budget Items: [Costs, expenses, or financial allocations]
- Revenue Impact: [Income or financial benefits]
- Cost Implications: [Financial risks or obligations]

**Action Items & Requirements:**
- Immediate Actions: [What needs to be done right away]
- Approval Requirements: [Who needs to sign off or authorize]
- Follow-up Tasks: [Ongoing or future responsibilities]

**Compliance & Risk Assessment**

**Potential Risks:**
- [HIGH/MEDIUM/LOW] [Specific risks with brief explanations]

**Missing Information:**
- [Critical details that should be present but aren't]
- [Questions that need answers for complete understanding]

**Compliance Considerations:**
- [Regulatory requirements that may apply]
- [Legal or policy implications identified]

**Actionable Recommendations**

**Immediate Actions Required:**
- [High-priority items needing attention]

**Follow-up Actions:**
- [Medium-term tasks and responsibilities]

**Stakeholder Communications:**
- [Who needs to be informed and about what]

**DOCUMENT CONTENT TO ANALYZE:**
{extracted_text}
"""
async def call_gemini_with_retry(prompt: str, request_id: str) -> str:
    """Call Gemini API with retry logic and proper error handling"""
    
    api_key = os.environ.get('GEMINI_API_KEY')
    if not api_key:
        logger.error(f"[{request_id}] ‚ùå GEMINI_API_KEY not found in environment")
        raise HTTPException(status_code=500, detail="Gemini API key not configured")
    
    # Gemini API endpoint
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash-latest:generateContent?key={api_key}"
    
    # Prepare the request payload
    payload = {
        "contents": [
            {
                "parts": [
                    {
                        "text": prompt
                    }
                ]
            }
        ],
        "generationConfig": {
            "temperature": 0.1,
            "topK": 40,
            "topP": 0.95,
            "maxOutputTokens": 8192,
            "stopSequences": []
        },
        "safetySettings": [
            {
                "category": "HARM_CATEGORY_HARASSMENT",
                "threshold": "BLOCK_MEDIUM_AND_ABOVE"
            },
            {
                "category": "HARM_CATEGORY_HATE_SPEECH",
                "threshold": "BLOCK_MEDIUM_AND_ABOVE"
            },
            {
                "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
                "threshold": "BLOCK_MEDIUM_AND_ABOVE"
            },
            {
                "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
                "threshold": "BLOCK_MEDIUM_AND_ABOVE"
            }
        ]
    }
    
    headers = {
        "Content-Type": "application/json",
        "User-Agent": "Document-Analysis-Middleware/2.0.0"
    }
    
    # Retry logic
    for attempt in range(MAX_RETRIES):
        try:
            logger.info(f"[{request_id}] üîÑ Calling Gemini API (attempt {attempt + 1}/{MAX_RETRIES})")
            
            response = requests_session.post(
                url,
                json=payload,
                headers=headers,
                timeout=GEMINI_TIMEOUT
            )
            
            # Log response details for debugging
            logger.info(f"[{request_id}] üìä Gemini API Response Status: {response.status_code}")
            
            if response.status_code == 200:
                result = response.json()
                
                # Extract the generated text
                try:
                    generated_text = result['candidates'][0]['content']['parts'][0]['text']
                    logger.info(f"[{request_id}] ‚úÖ Gemini API call successful, response length: {len(generated_text)} characters")
                    return generated_text.strip()
                except (KeyError, IndexError) as e:
                    logger.error(f"[{request_id}] ‚ùå Unexpected Gemini API response structure: {str(e)}")
                    logger.error(f"[{request_id}] Response: {result}")
                    raise HTTPException(status_code=500, detail="Unexpected API response format")
            
            elif response.status_code == 400:
                error_detail = extract_gemini_error(response, request_id)
                logger.error(f"[{request_id}] ‚ùå Gemini API bad request: {error_detail}")
                raise HTTPException(status_code=400, detail=f"Invalid request to Gemini API: {error_detail}")
            
            elif response.status_code == 401:
                logger.error(f"[{request_id}] ‚ùå Gemini API authentication failed")
                raise HTTPException(status_code=500, detail="Gemini API authentication failed")
            
            elif response.status_code == 403:
                error_detail = extract_gemini_error(response, request_id)
                logger.error(f"[{request_id}] ‚ùå Gemini API access denied: {error_detail}")
                raise HTTPException(status_code=403, detail=f"Access denied by Gemini API: {error_detail}")
            
            elif response.status_code == 429:
                # Rate limiting - retry with exponential backoff
                retry_delay = INITIAL_RETRY_DELAY * (2 ** attempt)
                logger.warning(f"[{request_id}] ‚ö†Ô∏è Rate limited by Gemini API, retrying in {retry_delay}s")
                
                if attempt < MAX_RETRIES - 1:  # Don't sleep on the last attempt
                    await asyncio.sleep(retry_delay)
                    continue
                else:
                    raise HTTPException(status_code=429, detail="Gemini API rate limit exceeded")
            
            elif response.status_code >= 500:
                # Server error - retry with exponential backoff
                retry_delay = INITIAL_RETRY_DELAY * (2 ** attempt)
                logger.warning(f"[{request_id}] ‚ö†Ô∏è Gemini API server error ({response.status_code}), retrying in {retry_delay}s")
                
                if attempt < MAX_RETRIES - 1:  # Don't sleep on the last attempt
                    await asyncio.sleep(retry_delay)
                    continue
                else:
                    error_detail = extract_gemini_error(response, request_id)
                    raise HTTPException(status_code=502, detail=f"Gemini API server error: {error_detail}")
            
            else:
                # Unexpected status code
                error_detail = extract_gemini_error(response, request_id)
                logger.error(f"[{request_id}] ‚ùå Unexpected Gemini API status {response.status_code}: {error_detail}")
                raise HTTPException(status_code=502, detail=f"Unexpected Gemini API response: {error_detail}")
                
        except requests.exceptions.Timeout:
            logger.warning(f"[{request_id}] ‚ö†Ô∏è Gemini API timeout (attempt {attempt + 1}/{MAX_RETRIES})")
            if attempt == MAX_RETRIES - 1:
                raise HTTPException(status_code=504, detail="Gemini API request timed out")
            await asyncio.sleep(INITIAL_RETRY_DELAY * (2 ** attempt))
            
        except requests.exceptions.ConnectionError:
            logger.warning(f"[{request_id}] ‚ö†Ô∏è Gemini API connection error (attempt {attempt + 1}/{MAX_RETRIES})")
            if attempt == MAX_RETRIES - 1:
                raise HTTPException(status_code=502, detail="Unable to connect to Gemini API")
            await asyncio.sleep(INITIAL_RETRY_DELAY * (2 ** attempt))
            
        except requests.exceptions.RequestException as e:
            logger.error(f"[{request_id}] ‚ùå Gemini API request failed: {str(e)}")
            if attempt == MAX_RETRIES - 1:
                raise HTTPException(status_code=502, detail=f"Gemini API request failed: {str(e)}")
            await asyncio.sleep(INITIAL_RETRY_DELAY * (2 ** attempt))
    
    # This should not be reached, but just in case
    raise HTTPException(status_code=500, detail="Maximum retries exceeded for Gemini API")


def extract_gemini_error(response: requests.Response, request_id: str) -> str:
    """Extract error details from Gemini API response"""
    try:
        error_data = response.json()
        
        # Try to extract the error message
        if 'error' in error_data:
            error_info = error_data['error']
            if isinstance(error_info, dict):
                message = error_info.get('message', 'Unknown error')
                code = error_info.get('code', 'Unknown code')
                return f"{message} (Code: {code})"
            else:
                return str(error_info)
        
        return f"HTTP {response.status_code} - {response.reason}"
        
    except json.JSONDecodeError:
        logger.warning(f"[{request_id}] ‚ö†Ô∏è Could not parse error response as JSON")
        return f"HTTP {response.status_code} - {response.reason}"
    except Exception as e:
        logger.warning(f"[{request_id}] ‚ö†Ô∏è Error extracting error details: {str(e)}")
        return f"HTTP {response.status_code} - {response.reason}"


# ----------------------------
# Optional: Follow-up Chat Endpoint
# ----------------------------

@app.post("/followup")
async def followup_chat(payload: FollowupPayload):
    """Handle follow-up questions about previously analyzed documents"""
    request_id = str(int(time.time() * 1000))
    logger.info(f"[{request_id}] üí¨ Processing follow-up chat with {len(payload.messages)} messages")
    
    try:
        # Input validation
        if not payload.messages:
            raise HTTPException(status_code=400, detail="Messages array cannot be empty")
        
        if len(payload.messages) > 50:  # Reasonable limit for conversation history
            raise HTTPException(status_code=400, detail="Too many messages in conversation history")
        
        # Build conversation prompt for Gemini
        conversation_prompt = build_followup_prompt(payload.messages)
        
        # Call Gemini API
        response = await call_gemini_with_retry(conversation_prompt, request_id)
        
        logger.info(f"[{request_id}] ‚úÖ Follow-up chat completed successfully")
        return {
            "response": response,
            "status": "success",
            "request_id": request_id,
            "timestamp": time.time()
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[{request_id}] ‚ùå Follow-up chat error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Follow-up chat failed: {str(e)}")


def build_followup_prompt(messages: List[Message]) -> str:
    """Build a conversational prompt from message history"""
    
    prompt_parts = [
        "You are an expert document analysis assistant. A user is asking follow-up questions about a document that was previously analyzed.",
        "",
        "INSTRUCTIONS:",
        "- Provide helpful, accurate answers based on the conversation context",
        "- If you need to refer to specific document details, ask for clarification",
        "- Keep responses concise and focused on the user's question",
        "- If the question is outside the scope of document analysis, politely redirect",
        "",
        "CONVERSATION HISTORY:"
    ]
    
    # Add message history
    for i, message in enumerate(messages):
        role = "USER" if message.role.lower() in ['user', 'human'] else "ASSISTANT"
        prompt_parts.append(f"{role}: {message.content.strip()}")
    
    prompt_parts.extend([
        "",
        "Please provide a helpful response to the user's most recent message."
    ])
    
    return "\n".join(prompt_parts)
